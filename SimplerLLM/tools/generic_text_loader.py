import newspaper
import os
import PyPDF2
import docx
from youtube_transcript_api import YouTubeTranscriptApi
import re
from urllib.parse import urlparse
from pydantic import BaseModel
from typing import Optional



class TextDocument(BaseModel):
    file_size: Optional[int] = None
    word_count: int
    character_count: int
    content: str
    title: Optional[str] = None
    url_or_path: Optional[str] = None

def load_text(input_path_or_url):
    # Check if the input is a URL
    input_path_or_url = str.lower(input_path_or_url)
    if re.match(r'http[s]?://', input_path_or_url):
        # Process based on URL content
        if "youtube.com" in input_path_or_url or "youtu.be" in input_path_or_url:
            content = __read_youtube_video_transcript(input_path_or_url)
            file_size = len(content.encode('utf-8'))  # Size in bytes
            return TextDocument(
                word_count=len(content.split()),
                character_count=len(content),
                content=content,
                file_size = file_size,
                url_or_path=input_path_or_url
            )
        else:
            article = __read_blog_from_url(input_path_or_url)
            if article is not None:
                file_size = len(article.text.encode('utf-8'))  # Size in bytes
                return TextDocument(
                    word_count=len(article.text.split()),
                    character_count=len(article.text),
                    content=article.text,
                    title=article.title,
                    file_size=file_size,
                    url_or_path=input_path_or_url
                )
    else:
        try:
            # Process based on file extension
            file_ext = os.path.splitext(input_path_or_url)[1].lower()
            if file_ext in ['.txt']:
                file_size, num_words, num_chars, content = __read_text_file(input_path_or_url)
            elif file_ext in ['.docx']:
                file_size, num_words, num_chars, content = __read_docx_file(input_path_or_url)
            elif file_ext in ['.pdf']:
                file_size, num_words, num_chars, content = __read_pdf_file(input_path_or_url)
            else:
                # Fallback: try reading as a text file
                file_size, num_words, num_chars, content = __read_text_file(input_path_or_url)

            return TextDocument(
                    file_size=file_size,
                    word_count=num_words,
                    character_count=num_chars,
                    content=content,
                    url_or_path=input_path_or_url
                )
        except Exception as e:
            raise ValueError(f"Error processing file: {e}")

    raise ValueError("Unable to process the input")






def __read_text_file(file_path):
    with open(file_path, 'r',encoding='utf-8') as file:
        content = file.read()

    file_size = os.path.getsize(file_path)
    words = content.split()
    num_words = len(words)
    num_chars = len(content)

    return file_size, num_words, num_chars, content

def __read_docx_file(file_path):
    file_size = os.path.getsize(file_path)
    doc = docx.Document(file_path)
    content = "\n".join([para.text for para in doc.paragraphs])

    words = content.split()
    num_words = len(words)
    num_chars = len(content)

    return file_size, num_words, num_chars, content

def __read_pdf_file(file_path):
    file_size = os.path.getsize(file_path)

    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        content = "".join([reader.pages[i].extract_text() for i in range(len(reader.pages))])

    words = content.split()
    num_words = len(words)
    num_chars = len(content)

    return file_size, num_words, num_chars, content

def __read_blog_from_url(url):
    """
    Extracts the text content from a given URL using the newspaper package.

    Parameters:
    url (str): The URL of the article to extract text from.

    Returns:
    str: The text content of the article if extraction is successful, None otherwise.
    """
    try:
        article = newspaper.Article(url)
        article.download()

        if article.download_state == 2:
            article.parse()
            return article
        else:
            print(f"An error occurred while fetching the article")
            return None
    except newspaper.ArticleException as e:
        print(f"An error occurred while fetching the article: {e}")
        return None

def __read_youtube_video_transcript(video_url):
    """
    Fetches the transcript of a YouTube video given its URL.

    Parameters:
    video_url (str): The URL of the YouTube video.

    Returns:
    str: The transcript of the video if available, raises an error otherwise.
    """
    # Enhanced regex to handle different YouTube URL formats
    match = re.search(r"(?:youtube\.com/watch\?v=|youtu\.be/)([\w-]+)", video_url)
    if match:
        video_id = match.group(1)
    else:
        raise ValueError("Invalid YouTube URL")

    try:
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        transcript_text = " ".join([line["text"] for line in transcript])
        return transcript_text
    except Exception as e:
        raise f"An error occurred while fetching the transcript: {e}"