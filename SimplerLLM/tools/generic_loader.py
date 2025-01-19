import newspaper
import os
import PyPDF2
import docx
from youtube_transcript_api import YouTubeTranscriptApi
import re
from urllib.parse import urlparse
from pydantic import BaseModel
from typing import Optional


class TextDocument(BaseModel):
    file_size: Optional[int] = None
    word_count: int
    character_count: int
    content: str
    title: Optional[str] = None
    url_or_path: Optional[str] = None


def load_content(input_path_or_url):
    """
    Load content from a given input path or URL.

    This function handles the following types of input:
    - URLs: Supports blog articles.
    - Local files: Supports .txt, .csv, .docx, and .pdf file extensions.
    Args:
        input_path_or_url (str)

    """
    # Check if the input is a URL
    if re.match(r"http[s]?://", input_path_or_url):
        article = __read_blog_from_url(input_path_or_url)
        if article is not None:
            file_size = len(article.text.encode("utf-8"))  # Size in bytes
            return TextDocument(
                word_count=len(article.text.split()),
                character_count=len(article.text),
                content=article.text,
                title=article.title,
                file_size=file_size,
                url_or_path=input_path_or_url,
            )
    else:
        try:
            # Process based on file extension
            file_ext = os.path.splitext(input_path_or_url)[1].lower()
            if file_ext in [".txt", ".csv"]:
                file_size, num_words, num_chars, content = __read_text_file(
                    input_path_or_url
                )
            elif file_ext in [".docx"]:
                file_size, num_words, num_chars, content = __read_docx_file(
                    input_path_or_url
                )
            elif file_ext in [".pdf"]:
                file_size, num_words, num_chars, content = __read_pdf_file(
                    input_path_or_url
                )

            else:
                # Fallback: try reading as a text file
                file_size, num_words, num_chars, content = __read_text_file(
                    input_path_or_url
                )

            return TextDocument(
                file_size=file_size,
                word_count=num_words,
                character_count=num_chars,
                content=content,
                url_or_path=input_path_or_url,
            )
        except Exception as e:
            raise ValueError(f"Error processing file: {e}")

    raise ValueError("Unable to process the input")


def __read_text_file(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        content = file.read()

    file_size = os.path.getsize(file_path)
    words = content.split()
    num_words = len(words)
    num_chars = len(content)

    return file_size, num_words, num_chars, content


def __read_docx_file(file_path):
    file_size = os.path.getsize(file_path)
    doc = docx.Document(file_path)
    content = "\n".join([para.text for para in doc.paragraphs])

    words = content.split()
    num_words = len(words)
    num_chars = len(content)

    return file_size, num_words, num_chars, content


def __read_pdf_file(file_path):
    file_size = os.path.getsize(file_path)

    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        content = "".join(
            [reader.pages[i].extract_text() for i in range(len(reader.pages))]
        )

    words = content.split()
    num_words = len(words)
    num_chars = len(content)

    return file_size, num_words, num_chars, content


def __read_blog_from_url(url):
    """
    Extracts the text content from a given URL using the newspaper package.

    Parameters:
    url (str): The URL of the article to extract text from.

    Returns:
    str: The text content of the article if extraction is successful, None otherwise.
    """
    try:
        article = newspaper.Article(url)
        article.download()

        if article.download_state == 2:
            article.parse()
            return article
        else:
            print("An error occurred while fetching the article")
            return None
    except newspaper.ArticleException as e:
        print(f"An error occurred while fetching the article: {e}")
        return None